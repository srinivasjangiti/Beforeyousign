"""
generate_docs.py
Run with: py scripts/generate_docs.py
Generates BeforeYouSign_Documentation.docx in the project root.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), '..', 'BeforeYouSign_Documentation.docx')

# ─── Color palette ─────────────────────────────────────────────────────────
C_BLACK  = RGBColor(0x1c, 0x19, 0x17)   # stone-900
C_GREY   = RGBColor(0x57, 0x53, 0x4e)   # stone-600
C_LIGHT  = RGBColor(0xe7, 0xe5, 0xe4)   # stone-200
C_AMBER  = RGBColor(0xb4, 0x53, 0x09)   # amber accent
C_WHITE  = RGBColor(0xff, 0xff, 0xff)
C_CODE   = RGBColor(0x1a, 0x1a, 0x1a)
C_CODEBG = RGBColor(0xf5, 0xf4, 0xf3)

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hr(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E7E5E4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def style_normal(run, size=11, color=C_GREY, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'

# ─── Document setup ─────────────────────────────────────────────────────────
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Heading styles
for lvl, sz, clr in [
    ('Heading 1', 26, C_BLACK),
    ('Heading 2', 18, C_BLACK),
    ('Heading 3', 14, C_AMBER),
]:
    s = doc.styles[lvl]
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = clr
    s.font.name  = 'Calibri'

Normal = doc.styles['Normal']
Normal.font.name = 'Calibri'
Normal.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, color=C_GREY, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_normal(run, size=size, color=color, bold=bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    style_normal(run, color=C_GREY)
    return p

def code_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = C_CODE
    # light shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F2F1')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def key_val(key, val):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    style_normal(r1, bold=True, color=C_BLACK)
    r2 = p.add_run(val)
    style_normal(r2, color=C_GREY)
    return p

def make_table(headers, rows):
    col_count = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size  = Pt(10)
        run.font.name  = 'Calibri'
        set_cell_bg(cell, '1C1917')

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        bg = 'FAFAF9' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.color.rgb = C_GREY
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            set_cell_bg(cell, bg)

    doc.add_paragraph()  # spacing after table
    return t

def page_break():
    doc.add_page_break()

def divider():
    add_hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# ── TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
r = tp.add_run("BeforeYouSign")
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = C_BLACK
r.font.name = 'Calibri'

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Complete Technical Documentation")
r2.font.size = Pt(20)
r2.font.color.rgb = C_GREY
r2.font.name = 'Calibri'

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("AI-Powered Contract Analysis Platform — Full Codebase Walkthrough")
r3.font.size = Pt(13)
r3.font.color.rgb = C_AMBER
r3.italic = True
r3.font.name = 'Calibri'

doc.add_paragraph()
divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
para(
    "BeforeYouSign is a Next.js 16 web application that gives everyday people "
    "institutional-grade contract intelligence. Users upload a PDF, DOCX, or TXT "
    "contract and receive an AI-powered risk assessment — including clause-by-clause "
    "analysis, red flags, negotiation strategies, and plain-English summaries."
)

h2("1.1 Mission")
para(
    '"Legal comprehension has never been democratized — until now." The platform aims to '
    "level the playing field between parties who have institutional legal resources and "
    "those who do not. Every contract you sign without full comprehension is a blind risk."
)

h2("1.2 Technology Stack")
make_table(
    ["Layer", "Technology", "Version"],
    [
        ["Framework",       "Next.js (App Router)",        "16.0.7"],
        ["Language",        "TypeScript",                   "5.x"],
        ["Styling",         "Tailwind CSS (v4)",            "4.x"],
        ["AI / LLM",        "NVIDIA NIM (OpenAI-compat.)",  "API v1"],
        ["Auth",            "NextAuth.js (v5 beta)",        "5.0.0-beta.30"],
        ["Database (MVP)",  "localStorage / Supabase",     "2.89"],
        ["PDF Parsing",     "pdf-parse",                    "2.4.5"],
        ["DOCX Parsing",    "mammoth",                      "1.11.0"],
        ["Charts",          "Recharts",                     "2.15.0"],
        ["Deployment",      "Vercel",                       "Serverless"],
    ]
)

h2("1.3 Key NPM Dependencies")
deps = [
    ("openai",            "OpenAI-compatible SDK used to call the NVIDIA NIM API"),
    ("pdf-parse",         "Extracts raw text from PDF binary buffers"),
    ("mammoth",           "Extracts raw text from Word (.docx) binary buffers"),
    ("next-auth",         "Full-featured authentication — OAuth + email/password"),
    ("bcryptjs",          "Password hashing for credential sign-in"),
    ("docx",              "Generates Word documents (used in export features)"),
    ("jspdf + html2canvas","Client-side PDF export"),
    ("recharts",          "React charting library used in dashboards"),
    ("@supabase/supabase-js","Optional cloud database integration"),
    ("lucide-react",      "Icon set used throughout the UI"),
    ("diff-match-patch",  "Text diff used in contract version comparison"),
]
for name, desc in deps:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{name}")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = C_BLACK
    r2 = p.add_run(f" — {desc}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_GREY

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. DIRECTORY STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Directory Structure")

tree_lines = [
    "beforeyousign/",
    "├── app/                Next.js App Router pages & API routes",
    "│   ├── api/            14 API route handlers (analyze, negotiate, drafting…)",
    "│   ├── analyze/        /analyze page — the core feature",
    "│   ├── dashboard/      User activity dashboard",
    "│   ├── negotiate/      AI negotiation assistant",
    "│   ├── templates/      Contract template browser",
    "│   └── … 35 more route folders",
    "├── components/         44 shared React components",
    "├── lib/                57 business logic & utility modules",
    "├── prisma/             Prisma ORM schema (future PostgreSQL)",
    "├── supabase/           Supabase SQL migration files",
    "├── public/             Static assets (icons, images)",
    "├── auth.ts             NextAuth.js configuration (OAuth + credentials)",
    "├── vercel.json         Serverless function timeout overrides",
    "├── next.config.ts      Next.js compiler options",
    "└── .env.example        Environment variable template",
]
for line in tree_lines:
    code_line(line)

doc.add_paragraph()

h2("2.1 app/ — Pages and API Routes")
para(
    "Next.js 16 App Router maps every folder under app/ to a URL. Files named "
    "page.tsx export the React UI component for that route. Files named route.ts "
    "export HTTP handler functions (GET, POST, etc.) that run server-side."
)
make_table(
    ["Route (URL)", "File", "Purpose"],
    [
        ["/",               "app/page.tsx",              "Marketing homepage"],
        ["/analyze",        "app/analyze/page.tsx",      "Contract upload + live AI analysis"],
        ["/dashboard",      "app/dashboard/page.tsx",    "User activity & history"],
        ["/negotiate",      "app/negotiate/page.tsx",    "AI negotiation coach"],
        ["/templates",      "app/templates/page.tsx",    "Contract template browser"],
        ["/clauses",        "app/clauses/page.tsx",      "5,000+ clause library"],
        ["/risk",           "app/risk/page.tsx",         "ML dispute probability prediction"],
        ["/compliance",     "app/compliance/page.tsx",   "Regulatory compliance scanner"],
        ["/blockchain",     "app/blockchain/page.tsx",   "Blockchain contract registry"],
        ["/voice",          "app/voice/page.tsx",        "Voice-to-contract creator"],
        ["/lawyers",        "app/lawyers/page.tsx",      "Vetted lawyer marketplace"],
        ["/esignature",     "app/esignature/page.tsx",   "Digital signature workflow"],
        ["POST /api/analyze","app/api/analyze/route.ts", "Core AI analysis (SSE streaming)"],
        ["POST /api/negotiate","app/api/negotiate/route.ts","Negotiation AI endpoint"],
        ["POST /api/drafting","app/api/drafting/route.ts","Contract drafting endpoint"],
        ["GET  /api/templates","app/api/templates/route.ts","Template CRUD operations"],
        ["POST /api/pdf",   "app/api/pdf/route.ts",      "PDF tools (merge, compress, etc.)"],
        ["POST /api/share",  "app/api/share/route.ts",   "Generate shareable links"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. CORE ANALYSIS PIPELINE
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Core Analysis Pipeline")
para(
    "The contract analysis workflow is the heart of BeforeYouSign. It follows a clean "
    "five-step pipeline — from file upload to rendered results — using Server-Sent Events "
    "(SSE) for real-time streaming so the user watches the AI think live."
)

h2("3.1 End-to-End Flow")
steps = [
    "1. User drops a contract file on the FileUpload component",
    "2. Browser POSTs multipart/form-data to /api/analyze",
    "3. Server validates the file, parses text, and builds an AI prompt",
    "4. Server opens an SSE stream and forwards NVIDIA NIM token-by-token deltas to the browser",
    "5. Browser reads the stream and updates UI in real time",
    "6. On stream end, full structured JSON is sent as a 'done' event and rendered",
]
for s in steps:
    bullet(s)

h2("3.2 File: app/api/analyze/route.ts — The Streaming API Handler")
para(
    "This is the most critical server-side file. It is a Next.js Route Handler "
    "that streams results using the Server-Sent Events (SSE) protocol. "
    "Two key module-level settings are:"
)
key_val("export const runtime", '"nodejs"  — runs in Node.js, not Vercel Edge, to support pdf-parse and mammoth')
key_val("export const maxDuration", '60  — gives this function a full 60-second execution budget')

h3("How it works, step by step")
sub_steps = [
    "Receives multipart/form-data via request.formData(); extracts the File object and jurisdiction string",
    "Calls validateContractFile(file) from lib/security.ts — checks size ≤ 10 MB, MIME type allowlist, filename safety",
    "Calls DocumentParser.validateFileType(file) — secondary MIME check",
    "Calls DocumentParser.parse(file) — converts PDF/DOCX/TXT binary to a plain text string",
    "Trims text to ANALYZE_MAX_PROMPT_CHARS (env var, default 6000) to stay within model token budget",
    "Calls ContractAnalyzer.buildAnalysisPrompt(text, jurisdiction) to produce the detailed JSON instruction prompt",
    "Creates an NVIDIA NIM streaming chat completion via createNvidiaClient()",
    "Iterates the async stream: each delta token is forwarded to the browser as an SSE 'chunk' event in real time",
    "After the stream ends, calls parseJsonResponse() on the accumulated text to get clean JSON",
    "Calls ContractAnalyzer.formatAnalysis() to normalise the JSON onto the TypeScript ContractAnalysis type",
    "Sends a 'done' SSE event with the full analysis object — the browser renders the results",
    "Any exception at any step sends an 'error' SSE event and closes the stream gracefully",
]
for s in sub_steps:
    bullet(s)

h3("SSE message types")
make_table(
    ["Type", "Payload fields", "Meaning"],
    [
        ["status", "message: string",             "Progress text shown below the spinner"],
        ["chunk",  "content: string",             "One raw AI text delta — appended to the terminal display"],
        ["done",   "analysis: ContractAnalysis",  "Complete, structured analysis — triggers result render"],
        ["error",  "error: string, requestId",    "Something failed — shown as an error card with Restart button"],
    ]
)

h2("3.3 File: lib/document-parser.ts — File Text Extraction")
para(
    "DocumentParser is a static class with one public method: parse(file). "
    "It routes to the correct private parser based on the file's MIME type."
)
bullet("PDF  (application/pdf)  → parsePDF(): lazy-imports pdf-parse dynamically to avoid Vercel cold-start crashes, then returns data.text")
bullet("DOCX (.docx)            → parseDOCX(): uses mammoth.extractRawText({ buffer }) to get clean plain text")
bullet("DOC  (.doc)             → also handled by parseDOCX() via application/msword MIME")
bullet("TXT  (text/plain)       → parseTXT(): converts Buffer to UTF-8 string directly")
bullet("validateFileType(file): checks MIME against allowlist [pdf, docx, doc, txt]")
bullet("validateFileSize(file, maxBytes): returns true if file.size ≤ maxBytes (default 10 MB)")

para(
    "⚠️  WHY LAZY IMPORT: pdf-parse does module-level file reads during require(). Vercel's "
    "serverless bundler places modules in a read-only virtual filesystem where these paths "
    "don't exist, causing an immediate crash at cold-start. Using dynamic import() inside "
    "the parsePDF() function means the code only runs when a PDF is actually being processed, "
    "not at module initialisation time.",
    color=C_AMBER, bold=True
)

h2("3.4 File: lib/contract-analyzer.ts — The AI Brain")
para("ContractAnalyzer is a pure-static class with four public members:")

h3("analyze() — non-streaming batch analysis")
batch_bullets = [
    "Used by advanced/non-streaming endpoints",
    "Trims contract text to ANALYZE_MAX_PROMPT_CHARS (default 6 000 chars) to stay within token limits",
    "Calls generateText() from nvidia-client for a single synchronous response",
    "Implements exponential-backoff retry: waits 750 ms → 1500 ms → … between attempts on HTTP 503/429 errors",
    "Throws descriptive errors for rate-limit (503/429), auth failure (401), and unknown errors",
    "Returns a fully typed ContractAnalysis object via formatAnalysis()",
]
for b in batch_bullets:
    bullet(b)

h3("buildAnalysisPrompt() — the master AI prompt")
para("Constructs the entire text sent to the LLM. Key design decisions:")
prompt_bullets = [
    "Jurisdiction-aware: parses 'US-California' style codes into 'California, United States' for the model",
    "Supports 60+ countries including US states, Canadian provinces, EU nations, Asia-Pacific, Middle East, Africa, Latin America",
    "Strict length constraints on every output field (e.g., plainLanguage max 100 chars) to keep the response small and fast",
    "JSON-only output instruction: model is told to return ONLY a JSON object — no markdown, no commentary",
    "Hard-coded schema in the prompt: ensures the model always returns the expected fields and types",
    "Focuses on 8 highest-risk clause types: IP transfers, liability caps, auto-renewals, termination restrictions, indemnification, non-compete, payment terms, arbitration",
]
for b in prompt_bullets:
    bullet(b)

h3("formatAnalysis() — normalising AI output")
para(
    "Maps the raw parsed JSON from the AI onto the strict ContractAnalysis TypeScript type. "
    "Every field has a safe fallback (empty arrays, 0 scores, default strings) so the UI "
    "never crashes even if the model skips a field or returns an unexpected type."
)

h3("Confidence scoring methods")
make_table(
    ["Method", "Baseline", "What reduces it"],
    [
        ["calculateOverallConfidence()", "85", "Short summary, empty clauses/recommendations"],
        ["calculateRiskConfidence()",    "88", "No flags despite high risk score; fewer than 5 clauses"],
        ["calculateClauseConfidence()",  "90", "Fewer than 30% of clauses have industryComparison data"],
    ]
)

h2("3.5 File: lib/nvidia-client.ts — Centralised AI Gateway")
para(
    "Every AI call in the app goes through this single module. The NVIDIA API key and base URL "
    "are configured here once and NEVER exposed to the browser."
)

h3("Model catalogue — NVIDIA_MODELS")
make_table(
    ["Key", "Model ID", "When used"],
    [
        ["primary",  "meta/llama-3.3-70b-instruct", "Highest capability; advanced features, negotiation, drafting"],
        ["fallback", "meta/llama-3.1-70b-instruct",  "Fallback when primary is busy"],
        ["fast",     "meta/llama-3.1-8b-instruct",   "Contract analysis (speed-optimised, 15–25 s vs 50–90 s)"],
    ]
)

h3("createNvidiaClient()")
para(
    "Reads NVIDIA_API_KEY from environment, creates an OpenAI SDK client pointed at "
    "https://integrate.api.nvidia.com/v1 — the NVIDIA NIM endpoint that is 100% "
    "compatible with the OpenAI Chat Completions API."
)

h3("generateText()")
para(
    "Single-prompt synchronous text generation. Uses AbortController to enforce a "
    "hard timeout (default 45 000 ms from NVIDIA_REQUEST_TIMEOUT_MS env var) so requests "
    "never hang indefinitely. Throws with a clear 'timed out' message on abort."
)

h3("generateWithSystem()")
para(
    "Like generateText() but accepts a separate system prompt. Used by the negotiation assistant "
    "and contract drafter where a role/persona needs to be set separately from the user question."
)

h3("parseJsonResponse<T>()")
para("Robust two-attempt JSON parser for AI model output:")
bullet("Attempt 1: strips markdown code fences (```json … ```) then JSON.parse()")
bullet("Attempt 2 (if attempt 1 fails): runs repairTruncatedJson() which walks the string character by character, tracks open brackets/strings, closes anything left open, then retries JSON.parse()")
bullet("This handles the common case where a model hits its token limit mid-JSON and the response is cut off abruptly")

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. FRONTEND PAGES
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Frontend Pages (app/)")

h2("4.1 Homepage — app/page.tsx")
para(
    "Pure marketing page rendered client-side ('use client'). No data-fetching. "
    "Uses Lucide icons and Tailwind CSS for a minimal dark/light aesthetic. Sections:"
)
hp_bullets = [
    "Hero: headline + CTA button linking to /analyze",
    "The Problem: explains the legal asymmetry problem with three stat cards",
    "The Solution: three feature grid cards (Risk Detection, Plain Language, Strategic Guidance)",
    "AI-Powered: 8-card grid showcasing advanced features (Negotiate, Risk, Benchmark, Voice, Blockchain, Clause Library, Obligation Tracker, Multi-Language)",
    "Social proof stats: 10,000+ Contracts Analysed, 5,000+ Protected Users, <30 s Average Time",
    "Footer with legal disclaimer",
]
for b in hp_bullets:
    bullet(b)

h2("4.2 Analyze Page — app/analyze/page.tsx  [THE CORE FEATURE]")
para(
    "The main user-facing page. A 'use client' React component that manages four UI states:"
)
states = [
    ("Upload state",    "Shows the FileUpload component and trust badges (Privacy Protected, Data Not Stored, AI model name)"),
    ("Analyzing state", "Shows a live terminal-style display of raw AI JSON output streaming in; progress messages above"),
    ("Error state",     "Shows the error message in a card with a Restart Analysis button"),
    ("Result state",    "Renders the full AnalysisResult component with all tabs and data"),
]
for name, desc in states:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r1.font.color.rgb = C_BLACK
    r2 = p.add_run(desc)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_GREY

h3("handleFileSelect() — SSE stream consumer")
stream_bullets = [
    "Posts multipart/form-data to /api/analyze using the browser's fetch() API",
    "Gets response.body.getReader() — a ReadableStream reader for incremental data",
    "Uses TextDecoder to decode Uint8Array chunks; buffers partial SSE frames across chunk boundaries",
    "Splits the buffer on '\\n\\n' (SSE event boundary) and finds lines starting with 'data: '",
    "Parses each data line as JSON and dispatches based on event type:",
    "  • 'status'  → setStatusMessage(data.message)",
    "  • 'chunk'   → setStreamingText(prev => prev + data.content)",
    "  • 'done'    → setAnalysis(data.analysis), hides loading state",
    "  • 'error'   → setError(data.error)",
    "Auto-scrolls the terminal div to the bottom on each new chunk via useEffect + ref",
]
for b in stream_bullets:
    bullet(b)

h2("4.3 Other Pages (brief summaries)")
make_table(
    ["Page", "Main Component", "What it does"],
    [
        ["Dashboard",   "UserDashboard.tsx",                 "Activity stats, past analyses, quick actions"],
        ["Negotiate",   "AIContractNegotiationAssistant.tsx","Chat-style AI coach with counter-proposals"],
        ["Templates",   "TemplatesLibrary.tsx",              "Browse & download pre-vetted templates"],
        ["Clauses",     "ClauseCard.tsx",                    "Search 5,000+ pre-vetted clauses by type/risk"],
        ["Risk",        "RiskGauge.tsx",                     "ML dispute-probability prediction tool"],
        ["Compliance",  "ComplianceMonitoring.tsx",          "GDPR/HIPAA/SOC2 compliance gap analysis"],
        ["Blockchain",  "BlockchainVerification.tsx",        "Cryptographic contract hash registry"],
        ["Voice",       "voice engine + ContractChat",       "Speech-to-contract via browser speech API"],
        ["Lawyers",     "LawyerMarketplace.tsx",             "Directory of vetted lawyers for consultation"],
        ["E-Signature", "ESignature.tsx",                    "Full digital signature workflow"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. REACT COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════
h1("5. React Components (components/)")

h2("5.1 FileUpload.tsx")
para(
    "Drag-and-drop upload widget powered by react-dropzone. Accepts PDF, DOCX, and TXT. "
    "On drop it validates the file type and size client-side before calling the parent's "
    "onFileSelect(file, jurisdiction) callback. Also renders a jurisdiction picker "
    "(country dropdown + optional state/province) that feeds into the AI prompt."
)

h2("5.2 AnalysisResult.tsx (~190 KB — the largest component)")
para("Renders the full contract analysis report. Major sub-sections:")
ar_bullets = [
    "Risk Score Gauge — circular SVG meter; colour-coded Green (0–39) / Amber (40–69) / Red (70–100)",
    "Document Metadata card — file name, type, parties, governing law, effective date, auto-renewal flag",
    "Executive Summary card — 2-sentence AI summary with confidence score",
    "Red Flags panel — each flag has severity badge (Warning / Danger / Critical), title, description, affected clauses, and recommendation",
    "Clause Cards — each clause shows: original text excerpt, plain-English explanation, risk level chip, fairness score bar, industry percentile, industry comparison, negotiation strategy (priority, leverage, approach, fallback positions, market precedents)",
    "Insights panel — missing clauses, inter-clause contradictions, unusual terms, strengths to preserve",
    "Recommendations — numbered actionable items",
    "Export bar — PDF (via jsPDF + html2canvas) and DOCX (via docx package) download buttons",
    "Confidence breakdown — per-dimension scores and model information",
]
for b in ar_bullets:
    bullet(b)

h2("5.3 Navbar.tsx (~45 KB)")
para(
    "Top navigation bar. Includes links to all major features, authentication state "
    "(sign-in button or user avatar dropdown), a mobile hamburger menu, and "
    "responsive collapsible sub-menus for the 40+ available routes. "
    "Reads NextAuth session to conditionally show user-specific controls."
)

h2("5.4 Other Notable Components")
make_table(
    ["Component", "Purpose"],
    [
        ["RiskGauge.tsx",                    "Animated circular risk-score dial with min/max thresholds"],
        ["ESignature.tsx",                   "Canvas-based signature drawing + typed name + audit trail"],
        ["AIContractDrafter.tsx",            "Wizard UI for describing and generating a contract draft"],
        ["AIContractNegotiationAssistant.tsx","Step-by-step negotiation coach with structured counter-proposals"],
        ["TeamCollaboration.tsx",            "Multi-user commenting, @mentions, and approval workflow"],
        ["ComplianceMonitoring.tsx",         "Compliance gap summary dashboard with regulation links"],
        ["BlockchainVerification.tsx",       "Generates and verifies SHA-256 / blockchain hashes for contracts"],
        ["ErrorBoundary.tsx",                "React error boundary — catches render errors, shows safe fallback"],
        ["ToastProvider.tsx",                "Global push-notification toast system (success/error/info)"],
        ["Loading.tsx",                      "Full-screen loading spinner for Suspense and page transitions"],
        ["Skeleton.tsx",                     "Placeholder shimmer UI for content loading states"],
        ["ProgressStepper.tsx",              "Multi-step wizard progress indicator"],
        ["BookingForm.tsx",                  "Lawyer consultation booking form with date/time picker"],
        ["PDFTools.tsx",                     "Merge, compress, and extract pages from PDF files"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. LIBRARY MODULES
# ═══════════════════════════════════════════════════════════════════════════
h1("6. Library Modules (lib/)")

h2("6.1 lib/types.ts — Core TypeScript Interfaces")
para("All data shapes that flow through the application live here.")
make_table(
    ["Interface / Type", "Purpose"],
    [
        ["ContractAnalysis",      "Top-level result type — holds every other sub-type"],
        ["ClauseAnalysis",        "One extracted clause: text, risk, plain language, fairness, industry comparison, negotiation strategy"],
        ["RedFlag",               "A critical risk found: type, severity, affected clause IDs, recommendation"],
        ["ContractInsights",      "Missing clauses, contradictions between clauses, unusual terms, strengths to keep"],
        ["ContractMetadata",      "File-level meta: name, size, parties, dates, governing law, document type, auto-renewal flag"],
        ["AnalysisConfidence",    "Model name/version, confidence scores, analysis date, notes"],
        ["IndustryBenchmark",     "Comparison of this contract vs. industry average per clause"],
        ["ComparativeInsight",    "Single comparative observation (stricter/fairer/standard/unusual)"],
        ["ClauseCategory (type)","Union: payment | termination | liability | intellectual_property | confidentiality | dispute_resolution | warranties | indemnification | non_compete | general | other"],
        ["RedFlagType (type)",    "Union: ip_transfer | unlimited_liability | auto_renewal | restricted_termination | one_sided_amendment | venue_forum | waiver_of_rights | confidentiality_overreach | indemnification | non_compete | payment_terms | dispute_resolution | other"],
    ]
)

h2("6.2 lib/security.ts — Security Utilities")
sec_items = [
    ("sanitizeInput(input)",         "HTML-encodes &, <, >, \", ', / to prevent XSS in all user-supplied text"),
    ("sanitizeHTML(html)",           "Strips script tags and on* event attributes while preserving a safe tag allowlist"),
    ("CSRFManager class",            "Generates/stores/validates CSRF tokens (32-char random, 24-hr TTL) in sessionStorage"),
    ("RateLimiter class",            "In-memory sliding window rate limiter; checkLimit(key, config) — max requests per window"),
    ("validateContractFile(file)",   "Validates file: 10 MB max, MIME allowlist, no null bytes, max 255-char filename"),
    ("SessionManager class",         "Creates/reads/refreshes/destroys user sessions (30-min TTL in sessionStorage)"),
    ("logSecurityEvent()",           "Writes audit log entries to localStorage; routes to external service in production"),
    ("validatePasswordStrength()",   "Scores 0-4 based on length, character variety, common-pattern detection"),
    ("generateSecureToken(length)",  "Web Crypto API generates cryptographically safe random hex tokens"),
    ("hashData(data)",               "SHA-256 hashes a string using SubtleCrypto for integrity verification"),
    ("CSP_HEADERS constant",         "Ready-to-use Content-Security-Policy + X-Frame-Options + X-Content-Type-Options headers"),
]
for name, desc in sec_items:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{name}: ")
    r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.color.rgb = C_BLACK
    r2 = p.add_run(desc)
    r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.color.rgb = C_GREY

h2("6.3 lib/database.ts — Data Persistence Layer (MVP)")
para(
    "An abstraction layer over localStorage, designed to be drop-in replaced by "
    "Supabase or PostgreSQL in production. Provides a typed CRUD interface:"
)
db_items = [
    ("getItem<T>(key)",              "Reads and JSON-parses a single value from localStorage"),
    ("setItem<T>(key, value)",       "JSON-serialises and writes; auto-runs cleanup() if QuotaExceededError is thrown"),
    ("getAllItems<T>(collection)",   "Reads a collection array, applies optional sorting and pagination"),
    ("insert<T>(collection, item)", "Appends an item; rejects duplicate IDs with a console warning"),
    ("update<T>(collection, id)",   "Merges partial updates into an existing item by ID"),
    ("delete<T>(collection, id)",   "Removes item from collection by ID"),
    ("find<T>(collection, pred)",   "Filters a collection with a custom predicate function"),
    ("In-memory cache layer",        "5-minute TTL cache; writes automatically invalidate relevant cache entries"),
    ("exportData() / importData()",  "Full JSON backup and restore of all localStorage data"),
    ("getStorageStats()",            "Reports used/available bytes and item count for the storage quota"),
    ("Schema versioning",            "On first init, migrate() creates indexes for faster filtering queries"),
]
for name, desc in db_items:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{name}: ")
    r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.color.rgb = C_BLACK
    r2 = p.add_run(desc)
    r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.color.rgb = C_GREY

h2("6.4 Other Notable Library Files")
make_table(
    ["File", "Purpose"],
    [
        ["advanced-analyzer.ts",           "Multi-step deep analysis: position tracking, cross-clause contradiction detection, missing clause identification"],
        ["ai-contract-drafter.ts",         "AI contract drafting: takes a description, returns a complete draft for NDA/employment/service/etc."],
        ["ai-negotiation-assistant.ts",    "Generates counter-proposals per clause with BATNA analysis using the 70B primary model"],
        ["smart-template-builder.ts 62KB", "Comprehensive template engine: variable merging, jurisdiction variants, completeness validation, pre-signing risk scoring"],
        ["contract-versioning.ts",         "Stores and retrieves contract revisions with diff metadata"],
        ["version-compare.ts",             "Side-by-side word-level diff using the diff-match-patch library"],
        ["compliance-security.ts",         "GDPR, HIPAA, SOC2 clause scanner — flags non-compliant provisions"],
        ["regulatory-compliance-scanner.ts","Matches clauses against known legal requirements per jurisdiction"],
        ["contract-lifecycle-automation.ts","Auto-extracts renewal dates and obligation deadlines; sends reminder notifications"],
        ["obligation-tracker.ts",          "Extracts and tracks all obligations with assignees and due dates"],
        ["ml-risk-predictor.ts",           "Heuristic ML model estimating dispute probability with 82-95% accuracy range"],
        ["multi-language-analyzer.ts",     "Routes non-English contracts through a translation layer before analysis"],
        ["voice-contract-engine.ts",       "Converts browser speech recognition transcripts to structured contract text"],
        ["export-manager.ts",              "Unified export: PDF (jsPDF), DOCX (docx), plain text, with formatting"],
        ["share-links.ts",                 "Generates expiring shareable links with cryptographic tokens"],
        ["lawyers-data.ts 21KB",           "Static data for the lawyer marketplace: profiles, specialisms, pricing, ratings"],
        ["templates-data.ts 124KB",        "Large library of full contract template text and metadata"],
        ["analytics.ts",                   "Privacy-preserving user event tracking (no PII stored)"],
        ["performance.ts",                 "Web Vitals (FCP, LCP, CLS) measurement and reporting utilities"],
        ["health.ts",                      "Service health check aggregator for uptime monitoring"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. AUTHENTICATION
# ═══════════════════════════════════════════════════════════════════════════
h1("7. Authentication — auth.ts")
para(
    "NextAuth v5 (beta) handles all authentication. The single file auth.ts is the "
    "source of truth for providers, session strategy, and callbacks."
)

h2("7.1 Supported Sign-In Methods")
bullet("GitHub OAuth — AUTH_GITHUB_ID + AUTH_GITHUB_SECRET environment variables")
bullet("Google OAuth — AUTH_GOOGLE_ID + AUTH_GOOGLE_SECRET environment variables")
bullet("Email + Password (Credentials) — passwords hashed with bcryptjs at cost factor 10")

h2("7.2 How Credentials Auth Works (step by step)")
cred_steps = [
    "User submits email + password on /auth/signin",
    "NextAuth calls the authorize() callback in auth.ts",
    "authorize() looks up the user in an in-memory Map by email (mock DB — replace with real DB in production)",
    "bcrypt.compare(submittedPassword, storedHash) validates the password",
    "On success, returns user object { id, name, email, image } which NextAuth stores in a JWT",
    "On failure, throws an Error with a descriptive message that NextAuth surfaces to the UI",
]
for s in cred_steps:
    bullet(s)

h2("7.3 JWT Session Strategy")
jwt_items = [
    ("strategy: 'jwt'",  "Sessions are stored in signed, encrypted HTTP-only cookies — no DB session table needed"),
    ("maxAge: 30 days",  "JWT tokens remain valid for 30 days before requiring re-login"),
    ("jwt() callback",   "Runs on every sign-in: injects user.id and account.provider into the token payload"),
    ("session() callback","Runs on every session read: copies token.id into session.user.id so components can read it"),
]
for name, desc in jwt_items:
    key_val(name, desc)

h2("7.4 Custom Auth Pages")
make_table(
    ["URL", "Purpose"],
    [
        ["/auth/signin",  "Custom sign-in form with email/password fields + GitHub/Google OAuth buttons"],
        ["/auth/signout", "Sign-out confirmation page"],
        ["/auth/error",   "Error display for failed or rejected auth attempts"],
        ["/auth/welcome", "Landing page shown to newly registered users"],
    ]
)

h2("7.5 OAuth Auto-Registration")
para(
    "The signIn() callback auto-creates a user record in the in-memory Map when a "
    "GitHub or Google user signs in for the first time. This means OAuth users can "
    "immediately use all features without a separate registration step."
)

h2("7.6 registerUser() helper")
para(
    "An exported async function that hashes a password with bcrypt and adds the user "
    "to the in-memory store. In production this should be replaced with a database "
    "INSERT operation plus email verification and a password-reset flow."
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════════════════
h1("8. Environment Variables")
para("Copy .env.example to .env.local and fill in the values listed below.")
make_table(
    ["Variable", "Required", "Description"],
    [
        ["NVIDIA_API_KEY",             "YES", "API key for NVIDIA NIM — get one at build.nvidia.com"],
        ["AUTH_SECRET",                "YES", "Random 32-byte secret for JWT signing (openssl rand -base64 32)"],
        ["NEXT_PUBLIC_APP_URL",        "YES", "Full public URL (e.g. https://app.example.com or http://localhost:3000)"],
        ["NEXT_PUBLIC_MAX_FILE_SIZE",  "No",  "Max upload in bytes (default: 10 485 760 = 10 MB)"],
        ["AUTH_GITHUB_ID",             "No",  "GitHub OAuth Client ID"],
        ["AUTH_GITHUB_SECRET",         "No",  "GitHub OAuth Client Secret"],
        ["AUTH_GOOGLE_ID",             "No",  "Google OAuth Client ID"],
        ["AUTH_GOOGLE_SECRET",         "No",  "Google OAuth Client Secret"],
        ["NEXT_PUBLIC_SUPABASE_URL",   "No",  "Supabase project URL (for cloud database)"],
        ["NEXT_PUBLIC_SUPABASE_ANON_KEY","No","Supabase anonymous/public API key"],
        ["SUPABASE_SERVICE_ROLE_KEY",  "No",  "Supabase admin key — server-only, never expose to browser"],
        ["ANALYZE_MAX_PROMPT_CHARS",   "No",  "Max contract chars sent to AI (default: 6 000)"],
        ["ANALYZE_MAX_OUTPUT_TOKENS",  "No",  "Max AI output tokens (default: 4 096)"],
        ["ANALYZE_MAX_RETRIES",        "No",  "Retry attempts on 503/429 rate-limit errors (default: 1)"],
        ["NVIDIA_REQUEST_TIMEOUT_MS",  "No",  "Hard timeout for NVIDIA API requests in ms (default: 45 000)"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════
h1("9. Deployment Configuration — vercel.json")
para(
    "The app deploys to Vercel serverless functions. vercel.json overrides the default "
    "10-second function timeout for the five AI-heavy endpoints that need more time:"
)
make_table(
    ["API Route", "maxDuration", "Why it needs more time"],
    [
        ["app/api/analyze/route.ts",          "60 s", "Contract parsing + NVIDIA NIM full-stream"],
        ["app/api/analyze-advanced/route.ts", "60 s", "Multi-pass deep analysis"],
        ["app/api/detect-clauses/route.ts",   "45 s", "Targeted clause detection pass"],
        ["app/api/negotiate/route.ts",        "60 s", "AI negotiation counter-proposal generation"],
        ["app/api/drafting/route.ts",         "60 s", "Full contract draft from scratch generation"],
    ]
)

h2("9.1 How SSE Streaming Solves the Timeout Problem")
para(
    "Previously the /api/analyze route waited for a complete JSON response (50–90 s) "
    "before returning anything — causing 504 Gateway Timeout errors. The fix:"
)
sse_bullets = [
    "An SSE 'status' event is sent immediately when the request arrives, opening the HTTP response stream",
    "Vercel's timeout is measured from the last byte sent, not the request start — so flushing data periodically keeps the connection alive",
    "The NVIDIA NIM stream sends token deltas continuously, so the connection is always active",
    "Total analysis time stays within the 60-second budget without any timeouts",
    "As a bonus, users see live AI output instead of a blank loading screen",
]
for b in sse_bullets:
    bullet(b)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. DATA FLOW DIAGRAMS
# ═══════════════════════════════════════════════════════════════════════════
h1("10. Data Flow — Annotated Diagrams")

h2("10.1 Contract Analysis Pipeline")
flow = [
    "User (Browser)",
    "  │",
    "  │  Drops contract file onto /analyze page",
    "  ▼",
    "FileUpload Component",
    "  │  onFileSelect(file, jurisdiction) callback triggered",
    "  ▼",
    "fetch POST /api/analyze  (multipart/form-data)",
    "  │",
    "  ▼",
    "app/api/analyze/route.ts  [Node.js runtime, 60 s budget]",
    "  │",
    "  ├─ 1. validateContractFile(file)  → security.ts checks",
    "  ├─ 2. DocumentParser.parse(file)  → plain text string",
    "  ├─ 3. trim to 6 000 chars         → speed optimisation",
    "  ├─ 4. buildAnalysisPrompt()       → JSON schema prompt",
    "  ├─ 5. createNvidiaClient()        → NVIDIA NIM connection",
    "  ├─ 6. stream chat completion      → token-by-token deltas",
    "  │      for each delta → SSE 'chunk' event to browser",
    "  ├─ 7. parseJsonResponse()         → clean JSON object",
    "  ├─ 8. formatAnalysis()            → typed ContractAnalysis",
    "  └─ 9. SSE 'done' event with full analysis",
    "  │",
    "  ▼",
    "Browser SSE reader  (analyze/page.tsx)",
    "  │  'status' → setStatusMessage()",
    "  │  'chunk'  → setStreamingText(prev => prev + content)",
    "  │  'done'   → setAnalysis(data.analysis)",
    "  │  'error'  → setError(data.error)",
    "  ▼",
    "AnalysisResult component renders the full multi-tabbed report",
]
for line in flow:
    code_line(line)

doc.add_paragraph()
h2("10.2 Credentials Authentication Flow")
auth_flow = [
    "User fills /auth/signin  (email + password)",
    "  │",
    "  ▼",
    "NextAuth.js — authorize() callback in auth.ts",
    "  │  1. Lookup user by email in in-memory Map",
    "  │  2. bcrypt.compare(password, storedHash)",
    "  │  3. Return user object { id, name, email }",
    "  ▼",
    "NextAuth creates JWT  (signed with AUTH_SECRET)",
    "  │  Payload: { id, name, email, provider, iat, exp }",
    "  ▼",
    "JWT stored in encrypted HTTP-only cookie  (30-day maxAge)",
    "  │",
    "  ▼",
    "session() callback exposes session.user.id to page components",
    "  │",
    "  ▼",
    "User is now authenticated — can access protected routes",
]
for line in auth_flow:
    code_line(line)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. SECURITY ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
h1("11. Security Architecture")

h2("11.1 File Upload Security (Defense in Depth)")
sec_bullets = [
    "MIME type allow-listing: both route.ts AND DocumentParser check against [pdf, docx, doc, txt]",
    "File size cap: 10 MB enforced before any parsing occurs — rejects before touching CPU-intensive parsers",
    "Null byte detection in filenames: prevents potential path traversal or injection via filename",
    "File name length cap: 255 characters maximum — prevents buffer overflow style attacks",
    "No persistence: files are parsed in memory and immediately garbage-collected — nothing written to disk or DB",
    "File content never reaches the browser: only the extracted text string is used server-side",
]
for b in sec_bullets:
    bullet(b)

h2("11.2 Input Sanitization")
bullet("All string inputs run through sanitizeInput() which HTML-encodes: & < > \" ' /")
bullet("This prevents reflected XSS where user input is rendered back into the DOM")
bullet("sanitizeHTML() provides safe HTML rendering for rich-text fields using an allowlist approach")

h2("11.3 API Key Security")
bullet("NVIDIA_API_KEY exists ONLY in server-side environment variables")
bullet("It is never in a NEXT_PUBLIC_ variable (which would be embedded in the client bundle)")
bullet("The createNvidiaClient() function is the single access point — easy to audit")
bullet("Vercel encrypts environment variables at rest and injects them only at function runtime")

h2("11.4 Authentication Security")
bullet("JWT signed with AUTH_SECRET (256-bit minimum recommended)")
bullet("Cookies are HTTP-only (not accessible from JavaScript) and encrypted")
bullet("CSRF protection via CSRFManager tokens stored in sessionStorage")
bullet("bcrypt cost factor 10 = ~100 ms per hash — makes brute-force attacks very slow")
bullet("Password strength validation with scoring and actionable feedback")

h2("11.5 Content Security Policy")
bullet("frame-ancestors: none — prevents clickjacking (iframe embedding)")
bullet("form-action: self — prevents cross-site form POST hijacking")
bullet("connect-src restricts all XHR/fetch to self and known trusted APIs")
bullet("script-src prevents loading of untrusted third-party scripts")

h2("11.6 Rate Limiting")
bullet("In-memory RateLimiter class provides per-key sliding window limiting")
bullet("Can be applied per IP, per user, or per API key in any route handler")
bullet("⚠️ Production note: in-memory limiting resets on each serverless function cold start — replace with Redis-backed (e.g., Upstash) for distributed enforcement")

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12. HOW TO RUN LOCALLY
# ═══════════════════════════════════════════════════════════════════════════
h1("12. How to Run Locally")

h2("Step 1 — Clone and Install")
code_line("git clone <repo-url>  &&  cd beforeyousign")
code_line("npm install")

h2("Step 2 — Create Environment File")
code_line("copy .env.example .env.local    # Windows")
code_line("cp  .env.example .env.local     # Mac / Linux")
para("Then edit .env.local and fill in:")
code_line("NVIDIA_API_KEY=nvapi-xxxxxxxxxxxxxxxxxxxx")
code_line("AUTH_SECRET=<run: openssl rand -base64 32>")
code_line("NEXT_PUBLIC_APP_URL=http://localhost:3000")

h2("Step 3 — Start Development Server")
code_line("npm run dev")
para("The app will be available at http://localhost:3000")

h2("Available npm Scripts")
make_table(
    ["Command", "What it does"],
    [
        ["npm run dev",   "Start Next.js development server with hot reload"],
        ["npm run build", "Build the production bundle (checks for TypeScript errors)"],
        ["npm run start", "Start the production server (after build)"],
        ["npm run lint",  "Run ESLint across the codebase"],
        ["npm run clean", "Delete .next cache + log files"],
        ["npm run fresh", "Clean + immediately restart dev server (full reset)"],
    ]
)

divider()
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 13. FAQ
# ═══════════════════════════════════════════════════════════════════════════
h1("13. Frequently Asked Questions")

faqs = [
    (
        "Q: Why NVIDIA NIM instead of OpenAI?",
        "NVIDIA NIM runs the same Meta Llama models on NVIDIA's GPU infrastructure with "
        "competitive latency and cost. The client is 100% OpenAI SDK-compatible — switching "
        "providers requires only changing the baseURL and model name in lib/nvidia-client.ts."
    ),
    (
        "Q: Is the contract stored anywhere?",
        "No. The file is parsed in memory on the server, the extracted text is sent to NVIDIA's "
        "API over HTTPS, and then immediately discarded. Nothing is written to disk or database. "
        "The analysis result may be cached in the user's browser localStorage, but that's under "
        "the user's own control and never sent to any server."
    ),
    (
        "Q: What happens if the AI response is cut off (token limit)?",
        "The repairTruncatedJson() function in lib/nvidia-client.ts handles this automatically. "
        "It walks the partial JSON string, detects unclosed brackets/strings, closes them "
        "syntactically, then retries JSON.parse(). This turns most truncated responses into "
        "a usable (if incomplete) ContractAnalysis object."
    ),
    (
        "Q: Why is pdf-parse imported with dynamic import() inside the function?",
        "pdf-parse executes module-level code that reads local files when it is require()-d. "
        "In Vercel's serverless environment (read-only bundled filesystem), these file paths "
        "don't exist, causing an instant crash at cold-start. Dynamic import() inside the "
        "parsePDF() function defers this to runtime — only when an actual PDF is being processed."
    ),
    (
        "Q: How does SSE streaming prevent 504 Gateway Timeout errors?",
        "Vercel's timeout applies to idle connections — if no bytes are sent for the timeout "
        "duration, the request fails. By sending a 'status' SSE event immediately (before any "
        "AI call), the HTTP response stream is opened and data flows from the first second. "
        "The NVIDIA stream then sends token deltas continuously, keeping the connection alive "
        "for the full 60-second budget."
    ),
    (
        "Q: What is the database.ts localStorage used for?",
        "It's the MVP persistence layer for browser-local data: saved analyses, bookmarks, "
        "audit logs, CSRF tokens, and sessions. In production it's intended to be replaced "
        "with Supabase or PostgreSQL via Prisma (schema already defined in prisma/)."
    ),
    (
        "Q: How do I add a new OAuth provider (e.g., Microsoft)?",
        "In auth.ts, import the provider from 'next-auth/providers/microsoft' and add it to "
        "the providers array with the appropriate clientId and clientSecret environment variables. "
        "NextAuth handles the OAuth flow automatically."
    ),
    (
        "Q: How do I change the AI model used for analysis?",
        "In lib/nvidia-client.ts, edit the NVIDIA_MODELS object. In app/api/analyze/route.ts, "
        "change NVIDIA_MODELS.fast to NVIDIA_MODELS.primary for higher quality (but slower) results. "
        "All other AI calls across the app must be updated individually if needed."
    ),
]
for q, a in faqs:
    h3(q)
    para(a)

divider()

# ─── Footer ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("© 2025 BeforeYouSign — Technical Documentation  •  Auto-generated from codebase analysis")
r.font.size   = Pt(9)
r.font.color.rgb = C_GREY
r.font.italic    = True
r.font.name      = "Calibri"

# ─── Save ───────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"[DONE] Documentation written to: {os.path.abspath(OUT)}")
